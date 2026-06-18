#!/usr/bin/env python3
"""
build_brief_docx.py - deterministic styled .docx builder for the pre-brief skill.

WHAT IT DOES
  Renders a one-page meeting brief as a styled Word .docx: a titled header,
  navy section headings, numbered points that each carry a [Kind, anchor] label,
  and a "Transcript anchors" section with bold Point-N labels and verbatim
  quotes. The styling lives in code, not in the model's per-run judgment, so
  every brief looks the same instead of relying on the Cowork Google Drive
  connector (which uploads plain text and cannot set fonts, color, or weight).

HOW THE SKILL USES IT
  1. Write the brief content (from Step 2) to a JSON file - see CONTENT SCHEMA below.
  2. Run:  python3 build_brief_docx.py content.json "Pre-Brief (Discovery) - Acme.docx"
  3. Deliver the resulting .docx.
  Needs python-docx (pip install python-docx if the runtime lacks it).

CUSTOMIZING (this is the part people change)
  - Colors / font / sizes: edit the CONFIG block right below. Hex without '#'.

CONTENT SCHEMA (the JSON the skill writes)
  {
    "title": "Pre-Brief (Discovery): Acme Co",
    "subtitle": "Discovery call",                       # optional one-liner under the title
    "sections": [
      {"heading": "What matters going in", "points": [
        {"n": 1, "kind": "Concern",          "anchor": "00:14:32",   "text": "Budget sign-off sits with the CFO, not on this call."},
        {"n": 2, "kind": "Research finding", "anchor": "prior brief", "text": "Closed a **Series B** in March 2026."}
      ]},
      {"heading": "Suggested agenda", "points": [
        {"n": 6, "kind": "Agenda item", "anchor": "00:31:10", "text": "Confirm the two-week pilot scope."}
      ]}
    ],
    "anchors": [
      {"n": 1, "anchor": "00:14:32",   "quote": "I'd have to run that past our CFO before we sign anything."},
      {"n": 2, "anchor": "prior brief", "quote": "Series B closed March 2026, 40M raised."}
    ]
  }
  Point numbering is continuous across sections; the anchors reference the same
  numbers. `anchor` is a timestamp for transcript-grounded points or a short
  source label (e.g. "prior brief", "proposal") for artifact-grounded ones.
  **bold** spans inside point text are honored.
"""

import json
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================ CONFIG (edit to customize) ============================
CONFIG = {
    "font": "Arial",            # document-wide font
    "navy": "1F4E79",           # section headings + the [Kind, anchor] label + Point-N label
    "title_color": "1A1A1A",    # brief title (near-black)
    "anchor_gray": "7F7F7F",    # anchor source label + verbatim quote
    "subtitle_gray": "7F7F7F",  # subtitle under the title
    "sizes": {                  # point sizes
        "title": 20, "subtitle": 12, "h1": 14, "point": 10.5, "anchor": 9.5, "label": 9,
    },
}
# ===================================================================================


def _rgb(hex_str):
    return RGBColor.from_string(hex_str)


def _set_base_font(doc):
    style = doc.styles["Normal"]
    style.font.name = CONFIG["font"]
    style.font.size = Pt(CONFIG["sizes"]["point"])
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), CONFIG["font"])


def _add_runs(paragraph, text, *, bold=False, italic=False, color=None, size=None):
    """Add text to a paragraph, honoring **bold** spans inside `text`."""
    parts = str(text).split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = paragraph.add_run(part)
        run.bold = bold or (i % 2 == 1)
        run.italic = italic
        if color:
            run.font.color.rgb = _rgb(color)
        if size:
            run.font.size = Pt(size)
    if not parts or all(p == "" for p in parts):  # empty text -> keep an empty run
        paragraph.add_run("")


def _title(doc, title, subtitle):
    s = CONFIG["sizes"]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _add_runs(p, title, bold=True, color=CONFIG["title_color"], size=s["title"])
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)
        _add_runs(sp, subtitle, italic=True, color=CONFIG["subtitle_gray"], size=s["subtitle"])


def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    _add_runs(p, text, bold=True, color=CONFIG["navy"], size=CONFIG["sizes"]["h1"])


def _point(doc, point):
    s = CONFIG["sizes"]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)  # hanging indent under the number
    _add_runs(p, f"{point.get('n')}. ", size=s["point"])
    kind = point.get("kind", "")
    anchor = point.get("anchor", "")
    label = f"[{kind}, {anchor}] " if anchor else f"[{kind}] "
    _add_runs(p, label, bold=True, color=CONFIG["navy"], size=s["label"])
    _add_runs(p, point.get("text", ""), size=s["point"])


def _anchor_line(doc, a):
    s = CONFIG["sizes"]
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _add_runs(p, f"Point {a.get('n')} ", bold=True, color=CONFIG["navy"], size=s["anchor"])
    if a.get("anchor"):
        _add_runs(p, f"[{a['anchor']}] ", color=CONFIG["anchor_gray"], size=s["anchor"])
    _add_runs(p, f"“{a.get('quote', '')}”", italic=True,
              color=CONFIG["anchor_gray"], size=s["anchor"])


def build(content, out_path):
    # Voice guard: em-dashes are banned in the brief. Fail loudly so they get rewritten,
    # rather than shipping them or blind-swapping to a hyphen.
    hits = []

    def _scan(v, where):
        if isinstance(v, str):
            if "—" in v:
                hits.append(f"{where}: {v[:70]}")
        elif isinstance(v, dict):
            for k, x in v.items():
                _scan(x, f"{where}.{k}")
        elif isinstance(v, list):
            for i, x in enumerate(v):
                _scan(x, f"{where}[{i}]")

    _scan(content, "content")
    if hits:
        raise ValueError(
            "Em-dashes (—) are not allowed in the brief. Rewrite each into separate "
            "sentences or use a comma, colon, or parentheses (never a hyphen), then rebuild:\n  "
            + "\n  ".join(hits)
        )

    # Cross-check: each point carries exactly one Transcript-anchors entry and vice versa
    # (1:1 on `n`). A silent mismatch produces anchors pointing at the wrong points, which a
    # reader cannot catch, so fail loudly here instead.
    point_ns = [pt.get("n") for sec in content.get("sections", []) for pt in sec.get("points", [])]
    anchor_ns = [a.get("n") for a in content.get("anchors", [])]
    orphan_anchors = sorted(set(anchor_ns) - set(point_ns), key=str)
    missing_anchors = sorted(set(point_ns) - set(anchor_ns), key=str)
    if orphan_anchors or missing_anchors:
        problems = []
        if orphan_anchors:
            problems.append(f"anchors with no matching point: {orphan_anchors}")
        if missing_anchors:
            problems.append(f"points with no Transcript-anchors entry: {missing_anchors}")
        raise ValueError(
            "Point and anchor numbers do not line up (each point needs exactly one matching "
            "anchor entry, and each anchor one matching point). Fix the numbering, then rebuild:\n  "
            + "\n  ".join(problems)
        )

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
    _set_base_font(doc)

    _title(doc, content.get("title", "Pre-Brief"), content.get("subtitle"))

    for sec in content.get("sections", []):
        _heading(doc, sec.get("heading", ""))
        for pt in sec.get("points", []):
            _point(doc, pt)

    anchors = content.get("anchors", [])
    if anchors:
        _heading(doc, "Transcript anchors")
        for a in anchors:
            _anchor_line(doc, a)

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 3:
        print("usage: python3 build_brief_docx.py <content.json> <out.docx>", file=sys.stderr)
        sys.exit(2)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        content = json.load(f)
    path = build(content, sys.argv[2])
    print(f"wrote {path}")


if __name__ == "__main__":
    main()
