#!/usr/bin/env python3
"""
build_docx.py - reusable, deterministic styled .docx builder for any skill.

WHY THIS EXISTS
  The Cowork Google Drive connector uploads plain text only: no fonts, color,
  weight, alignment, or table shading. To deliver a styled Word doc, the model
  writes structured JSON (content only) and this script renders the look. The
  styling lives in code, not in the model's per-run judgment, so every doc comes
  out identical instead of drifting run to run.

HOW TO REUSE IT (copy-in, not import)
  Skills ship as standalone ZIPs, so a shared import will not travel with the
  skill. Copy this file into your skill's `assets/` and adapt it:
    1. Edit the CONFIG block for your palette / font / sizes.
    2. Keep the generic blocks (h1-h3, p, bullets, table, status_table,
       numbered, signature, title_block, spacer). Add or delete specialized
       blocks as your doc needs - the render loop in build() is the one place
       to wire a new block type.
    3. Document your final block schema in this docstring so the model knows
       exactly what JSON to emit.

HOW THE SKILL USES IT
  1. Write the doc content to a JSON file - see CONTENT SCHEMA below.
  2. Run:  python3 build_docx.py content.json "Output Name.docx"
  3. Deliver the resulting .docx.
  Needs python-docx (pip install python-docx if the runtime lacks it).

CONTENT SCHEMA (the JSON the skill writes)
  {
    "title_block": {                                   # optional
      "eyebrow": "PROPOSAL",                           # optional small caps line
      "title": "Done-For-You Calling Engagement",
      "subtitle": "Completed-Conversations Program",   # optional italic line
      "columns": [                                     # optional N-column borderless row
        {"label": "PREPARED FOR", "lines": ["**Bridgepointe Advisors**", "Tony Lenci, Advisor"]},
        {"label": "PREPARED BY",  "lines": ["**RevCentric.ai**", "hunter@revcentric.ai"]}
      ],
      "footer": "May 21, 2026"                         # optional italic line (date etc.)
    },
    "blocks": [
      {"type": "h1", "text": "Executive Summary"},
      {"type": "h2", "text": "What We Sell"},
      {"type": "h3", "text": "Scope"},
      {"type": "p",  "text": "Plain text with **bold** spans supported."},
      {"type": "bullets", "items": ["First point", "Second **point**"]},
      {"type": "numbered", "items": [
        {"n": 1, "label": "Concern, 00:14:32", "text": "Budget sign-off sits with the CFO."},
        {"n": 2, "text": "Plain numbered item, no label."}
      ]},
      {"type": "table", "header": ["Conversations", "Meetings"],
       "rows": [["50", "5-7"]], "emphasis_rows": [0]},
      {"type": "status_table", "header": ["Disposition", "Status", "Definition"],
       "rows": [["Meeting Scheduled", "Billable", "..."], ["DNC", "Not Billable", "..."]],
       "status_col": 1},
      {"type": "signature", "parties": [
        {"title": "FOR THE COMPANY (Client)", "org": "TL Communications",
         "signer": "Tony Lenci", "role": "CEO"},
        {"title": "FOR THE SERVICE PROVIDER", "org": "RevCentric.ai",
         "sub": "A State of Iowa Company", "signer": "Hunter Deskin", "role": "Founder"}
      ]},
      {"type": "spacer"}
    ]
  }
  Block types: title_block (top-level key), h1, h2, h3, p, bullets, numbered,
  table, status_table, signature, spacer. **bold** spans are honored everywhere.
"""

import json
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================ CONFIG (edit to customize) ============================
# A clean navy / gray business palette. Hex without '#'. Swap these for brand colors.
CONFIG = {
    "font": "Arial",                 # document-wide font
    "navy": "1F4E79",                # headings + accent labels
    "title_color": "1A1A1A",         # main title (near-black)
    "label_gray": "7F7F7F",          # eyebrow, subtitle, column labels, quotes
    "table_header_fill": "D9E1F2",   # shaded table header row
    "table_emphasis_fill": "EFEFEF", # shaded total / emphasis row
    "billable_fill": "D9EAD3",       # green status cell
    "not_billable_fill": "F4CCCC",   # red/pink status cell
    "border_color": "BFBFBF",        # table grid line color
    "margins_in": {"left": 0.9, "right": 0.9, "top": 0.8, "bottom": 0.8},
    "sizes": {                       # point sizes
        "title": 26, "eyebrow": 10, "subtitle": 13,
        "h1": 16, "h2": 12.5, "h3": 11, "body": 10.5, "label": 9, "table": 9.5,
    },
}
# ===================================================================================


def _rgb(hex_str):
    return RGBColor.from_string(hex_str)


def _set_base_font(doc):
    style = doc.styles["Normal"]
    style.font.name = CONFIG["font"]
    style.font.size = Pt(CONFIG["sizes"]["body"])
    # ensure east-asian / cs fallbacks also use the font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), CONFIG["font"])


def _shade(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_borders(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")        # ~0.5pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def _no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tbl_pr.append(borders)


def _add_runs(paragraph, text, *, bold=False, italic=False, color=None, size=None, caps=False):
    """Add text to a paragraph, honoring **bold** spans inside `text`."""
    text = str(text)
    if caps:
        text = text.upper()
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = paragraph.add_run(part)
        run.bold = bold or (i % 2 == 1)
        run.italic = italic
        if color:
            run.font.color.rgb = _rgb(color)
        if size:
            run.font.size = Pt(size)
    if not parts or all(p == "" for p in parts):  # empty text -> keep an empty run
        paragraph.add_run("")


def _heading(doc, text, *, size, color, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True  # never strand a heading at a page bottom
    _add_runs(p, text, bold=True, color=color, size=size)
    return p


def _body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _add_runs(p, text, size=CONFIG["sizes"]["body"])
    return p


def _bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        _add_runs(p, item, size=CONFIG["sizes"]["body"])


def _numbered(doc, items):
    """Hanging-indent numbered items with an optional bracketed accent label."""
    s = CONFIG["sizes"]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)  # hang text under the number
        _add_runs(p, f"{item.get('n')}. ", size=s["body"])
        if item.get("label"):
            _add_runs(p, f"[{item['label']}] ", bold=True, color=CONFIG["navy"], size=s["label"])
        _add_runs(p, item.get("text", ""), size=s["body"])


def _data_table(doc, header, rows, emphasis_rows=None, status_col=None):
    emphasis_rows = set(emphasis_rows or [])
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # header
    for j, htext in enumerate(header):
        cell = table.rows[0].cells[j]
        _shade(cell, CONFIG["table_header_fill"])
        _set_cell_borders(cell, CONFIG["border_color"])
        cell.paragraphs[0].text = ""
        _add_runs(cell.paragraphs[0], htext, bold=True, size=CONFIG["sizes"]["table"])

    # body rows
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cell = cells[j]
            _set_cell_borders(cell, CONFIG["border_color"])
            if i in emphasis_rows:
                _shade(cell, CONFIG["table_emphasis_fill"])
            if status_col is not None and j == status_col:
                label = str(val).strip().lower()
                if label.startswith("billable") or label == "billable":
                    _shade(cell, CONFIG["billable_fill"])
                elif "not" in label:
                    _shade(cell, CONFIG["not_billable_fill"])
            cell.paragraphs[0].text = ""
            bold = i in emphasis_rows or (status_col is not None and j == status_col)
            _add_runs(cell.paragraphs[0], str(val), bold=bold, size=CONFIG["sizes"]["table"])
    return table


def _signature_block(doc, parties):
    table = doc.add_table(rows=1, cols=max(len(parties), 1))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    s = CONFIG["sizes"]

    def fill(cell, party):
        cell.paragraphs[0].text = ""
        _add_runs(cell.paragraphs[0], party.get("title", ""), bold=True,
                  color=CONFIG["label_gray"], size=s["label"], caps=True)
        org = cell.add_paragraph()
        org.paragraph_format.space_before = Pt(6)
        _add_runs(org, party.get("org", ""), bold=True, size=s["body"])
        if party.get("sub"):
            sub = cell.add_paragraph()
            _add_runs(sub, party["sub"], italic=True, color=CONFIG["label_gray"], size=s["label"])
        sig = cell.add_paragraph()
        sig.paragraph_format.space_before = Pt(14)
        _add_runs(sig, "Signature: ______________________________", size=s["body"])
        nm = cell.add_paragraph()
        _add_runs(nm, f"Name: {party.get('signer', '')}", size=s["body"])
        rl = cell.add_paragraph()
        _add_runs(rl, f"Title: {party.get('role', '')}", size=s["body"])
        dt = cell.add_paragraph()
        dt.paragraph_format.space_before = Pt(10)
        _add_runs(dt, "Date: ______________________", size=s["body"])

    for idx, party in enumerate(parties):
        fill(table.rows[0].cells[idx], party)
    for cell in table.rows[0].cells:
        _set_cell_borders(cell, CONFIG["border_color"])
    return table


def _title_block(doc, tb):
    s = CONFIG["sizes"]

    def centered():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        return p

    if tb.get("eyebrow"):
        _add_runs(centered(), tb["eyebrow"], bold=True, color=CONFIG["label_gray"],
                  size=s["eyebrow"], caps=True)
    p = centered()
    p.paragraph_format.space_after = Pt(3)
    _add_runs(p, tb.get("title", ""), bold=True, color=CONFIG["title_color"], size=s["title"])
    if tb.get("subtitle"):
        _add_runs(centered(), tb["subtitle"], italic=True, color=CONFIG["label_gray"], size=s["subtitle"])

    # N-column borderless row (e.g. prepared-for / prepared-by)
    columns = tb.get("columns", [])
    if columns:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        table = doc.add_table(rows=1, cols=len(columns))
        _no_table_borders(table)
        for idx, col in enumerate(columns):
            cell = table.rows[0].cells[idx]
            cell.paragraphs[0].text = ""
            _add_runs(cell.paragraphs[0], col.get("label", ""), bold=True,
                      color=CONFIG["label_gray"], size=s["label"], caps=True)
            for line in col.get("lines", []):
                lp = cell.add_paragraph()
                _add_runs(lp, line, size=s["body"])

    if tb.get("footer"):
        d = doc.add_paragraph()
        d.paragraph_format.space_before = Pt(8)
        _add_runs(d, tb["footer"], italic=True, color=CONFIG["label_gray"], size=s["body"])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _scan_em_dashes(content):
    """Voice guard: em-dashes are banned. Return a list of "where: snippet" hits."""
    hits = []

    def walk(v, where):
        if isinstance(v, str):
            if "—" in v:
                hits.append(f"{where}: {v[:70]}")
        elif isinstance(v, dict):
            for k, x in v.items():
                walk(x, f"{where}.{k}")
        elif isinstance(v, list):
            for i, x in enumerate(v):
                walk(x, f"{where}[{i}]")

    walk(content, "content")
    return hits


def build(content, out_path):
    hits = _scan_em_dashes(content)
    if hits:
        raise ValueError(
            "Em-dashes are not allowed. Rewrite each into separate sentences or use a "
            "comma, colon, or parentheses (never a hyphen), then rebuild:\n  "
            + "\n  ".join(hits)
        )

    doc = Document()
    m = CONFIG["margins_in"]
    for section in doc.sections:
        section.left_margin = Inches(m["left"])
        section.right_margin = Inches(m["right"])
        section.top_margin = Inches(m["top"])
        section.bottom_margin = Inches(m["bottom"])
    _set_base_font(doc)

    if content.get("title_block"):
        _title_block(doc, content["title_block"])

    s = CONFIG["sizes"]
    for block in content.get("blocks", []):
        btype = block.get("type")
        if btype == "h1":
            _heading(doc, block["text"], size=s["h1"], color=CONFIG["navy"], space_before=14)
        elif btype == "h2":
            _heading(doc, block["text"], size=s["h2"], color=CONFIG["navy"], space_before=10)
        elif btype == "h3":
            _heading(doc, block["text"], size=s["h3"], color=CONFIG["navy"], space_before=8)
        elif btype == "p":
            _body(doc, block["text"])
        elif btype == "bullets":
            _bullets(doc, block.get("items", []))
        elif btype == "numbered":
            _numbered(doc, block.get("items", []))
        elif btype == "table":
            _data_table(doc, block["header"], block.get("rows", []),
                        emphasis_rows=block.get("emphasis_rows"))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        elif btype == "status_table":
            _data_table(doc, block["header"], block.get("rows", []),
                        status_col=block.get("status_col"))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        elif btype == "signature":
            _signature_block(doc, block.get("parties", []))
        elif btype == "spacer":
            doc.add_paragraph()
        else:
            raise ValueError(f"unknown block type: {btype!r}")

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 3:
        print("usage: python3 build_docx.py <content.json> <out.docx>", file=sys.stderr)
        sys.exit(2)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        content = json.load(f)
    path = build(content, sys.argv[2])
    print(f"wrote {path}")


if __name__ == "__main__":
    main()
