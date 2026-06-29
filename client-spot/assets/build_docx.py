#!/usr/bin/env python3
"""
build_docx.py - deterministic styled .docx builder for the pre-call-prep skill.

WHY THIS EXISTS
  The Cowork Google Drive connector uploads plain text only: no fonts, color,
  weight, alignment, or table shading. To deliver a styled weekly check-in brief,
  the model writes structured JSON (content only) and this script renders the
  look. The styling lives in code, not in the model's per-run judgment, so every
  brief comes out identical instead of drifting run to run. This is the same
  proven pattern as the proposal builder and the SuperSDR pre-brief.

  CONFIG below is RevCentric's "Blue Professional" palette, orange-forward to match
  the disco-deck: orange section headings + eyebrow + accent rule, navy as the
  secondary (Sources labels, table headers), near-black ink title, muted gray
  locators. The brief and the deck read as one brand.

HOW THE SKILL USES IT
  1. Write the brief content to a JSON file - see CONTENT SCHEMA below.
  2. Run:  python3 build_docx.py content.json "RevCentric Weekly Check-in Brief - Acme (2026-06-17).docx"
  3. Deliver the resulting .docx (upload to the client's Drive folder for a View link).
  Needs python-docx (pip install python-docx if the runtime lacks it).

CONTENT SCHEMA (the JSON the skill writes - the six-section weekly-checkin brief)
  {
    "title_block": {
      "eyebrow": "WEEKLY CHECK-IN",                       # small orange-accent caps line
      "title": "RevCentric Weekly Check-in Brief: Acme Co",
      "columns": [                                         # borderless meta row
        {"label": "CLIENT",     "lines": ["**Acme Co**"]},
        {"label": "CHECK-IN",   "lines": ["2026-06-17"]},
        {"label": "PREPARED BY","lines": ["RevCentric.ai"]}
      ],
      "footer": "Read before the call. Every point is anchored in Sources."
    },
    "blocks": [
      {"type": "h1", "text": "Account status"},
      {"type": "p",  "text": "**At risk.** Booked meetings down to 4/wk from a 7/wk Q2 average; conversation rate steady at 6%."},

      {"type": "h1", "text": "What's working"},
      {"type": "bullets", "items": [
        "Connect rate up to 31% after the new call window (tracker, 2026-W24 summary)."
      ]},

      {"type": "h1", "text": "What's not working"},
      {"type": "bullets", "items": [
        "Meeting-set rate stalled two weeks running; most calls dying at gatekeeper."
      ]},

      {"type": "h1", "text": "Talking points"},
      {"type": "bullets", "items": [
        "Tighten the ICP to the two verticals converting best (SPOT: Target Accounts)."
      ]},

      {"type": "h1", "text": "Open action items"},
      {"type": "bullets", "items": [
        "Hunter to send the revised call script (last check-in, 00:21:40)."
      ]},

      {"type": "h1", "text": "Sources"},
      {"type": "numbered", "items": [
        {"n": 1, "label": "Master tracker", "text": "2026-W24 summary tab, booked-meeting trend and conversation rate."},
        {"n": 2, "label": "Slack",          "text": "#acme-revcentric > @hunter, 2026-06-14: \"calls keep dying at the gatekeeper\"."},
        {"n": 3, "label": "Prior check-in", "text": "00:21:40, Hunter commits to the revised script."},
        {"n": 4, "label": "SPOT",           "text": "Target Accounts section."}
      ]}
    ]
  }
  Use **bold** anywhere for emphasis (the Account status verdict, a key number).
  Block types: title_block (top-level key), h1, h2, h3, p, bullets, numbered,
  table, status_table, signature, spacer. Account status leads with a bold
  verdict; the middle sections are one-sentence bullets; Sources is a numbered
  list whose label is the source and whose text is the locator. Reorder `blocks`
  to reorder the doc, no code change. A source that was unavailable is still
  named (e.g. a bullet stating the prior transcript was unavailable), never
  silently dropped.
"""

import json
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================ CONFIG (edit to customize) ============================
# RevCentric "Blue Professional" palette, matched to the disco-deck template
# (deck.css :root). Hex without '#'.
CONFIG = {
    "font": "Arial",                 # document-wide font (portable; the deck's Space Grotesk/Inter are not on most hosts)
    "accent": "FC4E36",              # RevCentric orange (deck --orange): section headings, eyebrow, accent rule
    "navy": "1E4266",                # secondary: Sources labels, table headers, sub-headings (deck --navy)
    "title_color": "16243A",         # main title, near-black (deck --ink)
    "label_gray": "6B6B6B",          # subtitle, column labels, footer (deck --muted)
    "table_header_fill": "E4EAF1",   # shaded table header row (pale navy tint)
    "table_emphasis_fill": "F2F4F7", # shaded total / emphasis row
    "billable_fill": "D9EAD3",       # green status cell (healthy)
    "not_billable_fill": "F4CCCC",   # red/pink status cell (at risk)
    "border_color": "C9D2DD",        # table grid line (navy-tinted hairline, deck --line)
    "margins_in": {"left": 0.9, "right": 0.9, "top": 0.8, "bottom": 0.8},
    "sizes": {                       # point sizes
        "title": 24, "eyebrow": 10, "subtitle": 13,
        "h1": 15, "h2": 12.5, "h3": 11, "body": 10.5, "label": 9, "table": 9.5,
    },
}
# ===================================================================================


def _rgb(hex_str):
    return RGBColor.from_string(hex_str)


def _set_base_font(doc):
    style = doc.styles["Normal"]
    style.font.name = CONFIG["font"]
    style.font.size = Pt(CONFIG["sizes"]["body"])
    # ensure east-asian / cs fallbacks also use the font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), CONFIG["font"])


def _shade(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_borders(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")        # ~0.5pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def _accent_rule(doc, width_in=1.2):
    """A short, thick orange bar - the deck's signature accent-line under the headline."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    _no_table_borders(t)
    row = t.rows[0]
    row.height = Pt(3)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    cell.width = Inches(width_in)
    _shade(cell, CONFIG["accent"])
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(3)
    run = p.add_run("")
    run.font.size = Pt(2)
    return t


def _no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tbl_pr.append(borders)


def _add_runs(paragraph, text, *, bold=False, italic=False, color=None, size=None, caps=False):
    """Add text to a paragraph, honoring **bold** spans inside `text`."""
    text = str(text)
    if caps:
        text = text.upper()
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = paragraph.add_run(part)
        run.bold = bold or (i % 2 == 1)
        run.italic = italic
        if color:
            run.font.color.rgb = _rgb(color)
        if size:
            run.font.size = Pt(size)
    if not parts or all(p == "" for p in parts):  # empty text -> keep an empty run
        paragraph.add_run("")


def _heading(doc, text, *, size, color, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True  # never strand a heading at a page bottom
    _add_runs(p, text, bold=True, color=color, size=size)
    return p


def _body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _add_runs(p, text, size=CONFIG["sizes"]["body"])
    return p


def _bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        _add_runs(p, item, size=CONFIG["sizes"]["body"])


def _numbered(doc, items):
    """Hanging-indent numbered items with an optional bracketed accent label."""
    s = CONFIG["sizes"]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)  # hang text under the number
        _add_runs(p, f"{item.get('n')}. ", size=s["body"])
        if item.get("label"):
            _add_runs(p, f"[{item['label']}] ", bold=True, color=CONFIG["navy"], size=s["label"])
        _add_runs(p, item.get("text", ""), size=s["body"])


def _data_table(doc, header, rows, emphasis_rows=None, status_col=None):
    emphasis_rows = set(emphasis_rows or [])
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # header
    for j, htext in enumerate(header):
        cell = table.rows[0].cells[j]
        _shade(cell, CONFIG["table_header_fill"])
        _set_cell_borders(cell, CONFIG["border_color"])
        cell.paragraphs[0].text = ""
        _add_runs(cell.paragraphs[0], htext, bold=True, size=CONFIG["sizes"]["table"])

    # body rows
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cell = cells[j]
            _set_cell_borders(cell, CONFIG["border_color"])
            if i in emphasis_rows:
                _shade(cell, CONFIG["table_emphasis_fill"])
            if status_col is not None and j == status_col:
                label = str(val).strip().lower()
                # green: billable / healthy; red: not billable / at risk / needs attention
                if label == "healthy" or (label.startswith("billable") and "not" not in label):
                    _shade(cell, CONFIG["billable_fill"])
                elif "not" in label or "risk" in label or "attention" in label:
                    _shade(cell, CONFIG["not_billable_fill"])
            cell.paragraphs[0].text = ""
            bold = i in emphasis_rows or (status_col is not None and j == status_col)
            _add_runs(cell.paragraphs[0], str(val), bold=bold, size=CONFIG["sizes"]["table"])
    return table


def _signature_block(doc, parties):
    table = doc.add_table(rows=1, cols=max(len(parties), 1))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    s = CONFIG["sizes"]

    def fill(cell, party):
        cell.paragraphs[0].text = ""
        _add_runs(cell.paragraphs[0], party.get("title", ""), bold=True,
                  color=CONFIG["label_gray"], size=s["label"], caps=True)
        org = cell.add_paragraph()
        org.paragraph_format.space_before = Pt(6)
        _add_runs(org, party.get("org", ""), bold=True, size=s["body"])
        if party.get("sub"):
            sub = cell.add_paragraph()
            _add_runs(sub, party["sub"], italic=True, color=CONFIG["label_gray"], size=s["label"])
        sig = cell.add_paragraph()
        sig.paragraph_format.space_before = Pt(14)
        _add_runs(sig, "Signature: ______________________________", size=s["body"])
        nm = cell.add_paragraph()
        _add_runs(nm, f"Name: {party.get('signer', '')}", size=s["body"])
        rl = cell.add_paragraph()
        _add_runs(rl, f"Title: {party.get('role', '')}", size=s["body"])
        dt = cell.add_paragraph()
        dt.paragraph_format.space_before = Pt(10)
        _add_runs(dt, "Date: ______________________", size=s["body"])

    for idx, party in enumerate(parties):
        fill(table.rows[0].cells[idx], party)
    for cell in table.rows[0].cells:
        _set_cell_borders(cell, CONFIG["border_color"])
    return table


def _title_block(doc, tb):
    s = CONFIG["sizes"]

    def centered():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        return p

    if tb.get("eyebrow"):
        _add_runs(centered(), tb["eyebrow"], bold=True, color=CONFIG["accent"],
                  size=s["eyebrow"], caps=True)
    p = centered()
    p.paragraph_format.space_after = Pt(3)
    _add_runs(p, tb.get("title", ""), bold=True, color=CONFIG["title_color"], size=s["title"])
    if tb.get("subtitle"):
        _add_runs(centered(), tb["subtitle"], italic=True, color=CONFIG["label_gray"], size=s["subtitle"])
    _accent_rule(doc)  # orange brand bar under the title

    # N-column borderless row (e.g. client / check-in / prepared-by)
    columns = tb.get("columns", [])
    if columns:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        table = doc.add_table(rows=1, cols=len(columns))
        _no_table_borders(table)
        for idx, col in enumerate(columns):
            cell = table.rows[0].cells[idx]
            cell.paragraphs[0].text = ""
            _add_runs(cell.paragraphs[0], col.get("label", ""), bold=True,
                      color=CONFIG["label_gray"], size=s["label"], caps=True)
            for line in col.get("lines", []):
                lp = cell.add_paragraph()
                _add_runs(lp, line, size=s["body"])

    if tb.get("footer"):
        d = doc.add_paragraph()
        d.paragraph_format.space_before = Pt(8)
        _add_runs(d, tb["footer"], italic=True, color=CONFIG["label_gray"], size=s["body"])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _scan_em_dashes(content):
    """Voice guard: em-dashes are banned. Return a list of "where: snippet" hits."""
    hits = []

    def walk(v, where):
        if isinstance(v, str):
            if "—" in v:
                hits.append(f"{where}: {v[:70]}")
        elif isinstance(v, dict):
            for k, x in v.items():
                walk(x, f"{where}.{k}")
        elif isinstance(v, list):
            for i, x in enumerate(v):
                walk(x, f"{where}[{i}]")

    walk(content, "content")
    return hits


def build(content, out_path):
    hits = _scan_em_dashes(content)
    if hits:
        raise ValueError(
            "Em-dashes are not allowed. Rewrite each into separate sentences or use a "
            "comma, colon, or parentheses (never a hyphen), then rebuild:\n  "
            + "\n  ".join(hits)
        )

    doc = Document()
    m = CONFIG["margins_in"]
    for section in doc.sections:
        section.left_margin = Inches(m["left"])
        section.right_margin = Inches(m["right"])
        section.top_margin = Inches(m["top"])
        section.bottom_margin = Inches(m["bottom"])
    _set_base_font(doc)

    if content.get("title_block"):
        _title_block(doc, content["title_block"])

    s = CONFIG["sizes"]
    for block in content.get("blocks", []):
        btype = block.get("type")
        if btype == "h1":
            _heading(doc, block["text"], size=s["h1"], color=CONFIG["accent"], space_before=14)
        elif btype == "h2":
            _heading(doc, block["text"], size=s["h2"], color=CONFIG["navy"], space_before=10)
        elif btype == "h3":
            _heading(doc, block["text"], size=s["h3"], color=CONFIG["navy"], space_before=8)
        elif btype == "p":
            _body(doc, block["text"])
        elif btype == "bullets":
            _bullets(doc, block.get("items", []))
        elif btype == "numbered":
            _numbered(doc, block.get("items", []))
        elif btype == "table":
            _data_table(doc, block["header"], block.get("rows", []),
                        emphasis_rows=block.get("emphasis_rows"))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        elif btype == "status_table":
            _data_table(doc, block["header"], block.get("rows", []),
                        status_col=block.get("status_col"))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        elif btype == "signature":
            _signature_block(doc, block.get("parties", []))
        elif btype == "spacer":
            doc.add_paragraph()
        else:
            raise ValueError(f"unknown block type: {btype!r}")

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 3:
        print("usage: python3 build_docx.py <content.json> <out.docx>", file=sys.stderr)
        sys.exit(2)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        content = json.load(f)
    path = build(content, sys.argv[2])
    print(f"wrote {path}")


if __name__ == "__main__":
    main()
