#!/usr/bin/env python3
"""
build_proposal_docx.py - deterministic styled .docx builder for the Proposal Doc Builder skill.

WHAT IT DOES
  Renders a branded, send-ready proposal as a Word .docx that matches the
  RevCentric / Bridgepointe gold-standard layout: centered title block, navy
  headings, two-column "Prepared for / by", shaded table headers, green/red
  Appendix status cells, and a signature block. The look is guaranteed because
  the styling lives in code, not in the model's per-run judgment.

HOW THE SKILL USES IT
  1. Write the proposal content (from Step 4) to a JSON file - see CONTENT SCHEMA below.
  2. Run:  python3 build_proposal_docx.py content.json "Agency Proposal - Prospect.docx"
  3. Deliver the resulting .docx.
  Needs python-docx (pip install python-docx if the runtime lacks it).

CUSTOMIZING (this is the part people change)
  - Colors / font / sizes: edit the CONFIG block right below. Hex without '#'.
  - Section order / content: it follows the order of "blocks" in the JSON, so
    reordering or dropping blocks reorganizes the doc. No code change needed.

CONTENT SCHEMA (the JSON the skill writes)
  {
    "title_block": {
      "eyebrow": "PROPOSAL",
      "title": "Done-For-You Calling Engagement",
      "subtitle": "Completed-Conversations Outbound Program",   # optional
      "prepared_for": {"label": "PREPARED FOR", "org": "Bridgepointe Advisors", "person": "Tony Lenci, Advisor"},
      "prepared_by":  {"label": "PREPARED BY",  "org": "RevCentric.ai", "person": "Hunter Deskin - hunter@revcentric.ai"},
      "date": "May 21, 2026"
    },
    "blocks": [
      {"type": "h1", "text": "Executive Summary"},
      {"type": "h2", "text": "What Bridgepointe Sells"},
      {"type": "p",  "text": "Plain text with **bold** spans supported."},
      {"type": "bullets", "items": ["First point", "Second **point**"]},
      {"type": "table", "header": ["Conversations", "Meetings"], "rows": [["50", "5-7"]], "emphasis_rows": [0]},
      {"type": "status_table", "header": ["Disposition", "Status", "Definition"],
       "rows": [["Meeting Scheduled", "Billable", "..."], ["DNC", "Not Billable", "..."]], "status_col": 1},
      {"type": "signature",
       "company":  {"title": "FOR THE COMPANY (Client)", "org": "TL Communications", "signer": "Tony Lenci", "role": "CEO"},
       "provider": {"title": "FOR THE SERVICE PROVIDER", "org": "Deskin Enterprises DBA RevCentric.ai",
                    "sub": "A State of Iowa Company", "signer": "Hunter Deskin", "role": "Founder, RevCentric.ai"}}
    ]
  }
  Block types: h1, h2, h3, p, bullets, table, status_table, signature, spacer.
"""

import json
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================ CONFIG (edit to customize) ============================
CONFIG = {
    "font": "Arial",            # document-wide font
    "navy": "1F4E79",           # section / subsection heading color
    "title_color": "1A1A1A",    # main engagement title (near-black)
    "eyebrow_gray": "7F7F7F",   # "PROPOSAL" eyebrow + section labels
    "label_gray": "7F7F7F",     # "PREPARED FOR/BY" labels
    "table_header_fill": "D9E1F2",   # shaded table header row
    "table_emphasis_fill": "EFEFEF", # shaded total / emphasis row
    "billable_fill": "D9EAD3",       # green status cell
    "not_billable_fill": "F4CCCC",   # red/pink status cell
    "border_color": "BFBFBF",        # table grid line color
    "sizes": {                   # point sizes
        "title": 26, "eyebrow": 10, "subtitle": 13,
        "h1": 16, "h2": 12.5, "h3": 11, "body": 10.5, "label": 9, "table": 9.5,
    },
}
# ===================================================================================


def _rgb(hex_str):
    return RGBColor.from_string(hex_str)


def _set_base_font(doc):
    style = doc.styles["Normal"]
    style.font.name = CONFIG["font"]
    style.font.size = Pt(CONFIG["sizes"]["body"])
    # ensure east-asian / cs fallbacks also use the font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), CONFIG["font"])


def _shade(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_borders(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")        # ~0.5pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)


def _no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tbl_pr.append(borders)


def _add_runs(paragraph, text, *, bold=False, italic=False, color=None, size=None, caps=False):
    """Add text to a paragraph, honoring **bold** spans inside `text`."""
    if caps:
        text = text.upper()
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = paragraph.add_run(part)
        run.bold = bold or (i % 2 == 1)
        run.italic = italic
        if color:
            run.font.color.rgb = _rgb(color)
        if size:
            run.font.size = Pt(size)
    if not parts or all(p == "" for p in parts):  # empty text -> keep an empty run
        paragraph.add_run("")


def _heading(doc, text, *, size, color, space_before=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    _add_runs(p, text, bold=True, color=color, size=size)
    return p


def _body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _add_runs(p, text, size=CONFIG["sizes"]["body"])
    return p


def _bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        _add_runs(p, item, size=CONFIG["sizes"]["body"])


def _data_table(doc, header, rows, emphasis_rows=None, status_col=None):
    emphasis_rows = set(emphasis_rows or [])
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # header
    for j, htext in enumerate(header):
        cell = table.rows[0].cells[j]
        _shade(cell, CONFIG["table_header_fill"])
        _set_cell_borders(cell, CONFIG["border_color"])
        cell.paragraphs[0].text = ""
        _add_runs(cell.paragraphs[0], htext, bold=True, size=CONFIG["sizes"]["table"])

    # body rows
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cell = cells[j]
            _set_cell_borders(cell, CONFIG["border_color"])
            if i in emphasis_rows:
                _shade(cell, CONFIG["table_emphasis_fill"])
            if status_col is not None and j == status_col:
                label = str(val).strip().lower()
                if label.startswith("billable") or label == "billable":
                    _shade(cell, CONFIG["billable_fill"])
                elif "not" in label:
                    _shade(cell, CONFIG["not_billable_fill"])
            cell.paragraphs[0].text = ""
            bold = i in emphasis_rows or (status_col is not None and j == status_col)
            _add_runs(cell.paragraphs[0], str(val), bold=bold, size=CONFIG["sizes"]["table"])
    return table


def _signature_block(doc, company, provider):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def fill(cell, party):
        cell.paragraphs[0].text = ""
        _add_runs(cell.paragraphs[0], party.get("title", ""), bold=True,
                  color=CONFIG["eyebrow_gray"], size=CONFIG["sizes"]["label"], caps=True)
        org = cell.add_paragraph()
        org.paragraph_format.space_before = Pt(6)
        _add_runs(org, party.get("org", ""), bold=True, size=CONFIG["sizes"]["body"])
        if party.get("sub"):
            sub = cell.add_paragraph()
            _add_runs(sub, party["sub"], italic=True, color=CONFIG["eyebrow_gray"],
                      size=CONFIG["sizes"]["label"])
        sig = cell.add_paragraph()
        sig.paragraph_format.space_before = Pt(14)
        _add_runs(sig, "Signature: ______________________________", size=CONFIG["sizes"]["body"])
        nm = cell.add_paragraph()
        _add_runs(nm, f"Name: {party.get('signer', '')}", size=CONFIG["sizes"]["body"])
        rl = cell.add_paragraph()
        _add_runs(rl, f"Title: {party.get('role', '')}", size=CONFIG["sizes"]["body"])
        dt = cell.add_paragraph()
        dt.paragraph_format.space_before = Pt(10)
        _add_runs(dt, "Date: ______________________", size=CONFIG["sizes"]["body"])

    fill(table.rows[0].cells[0], company)
    fill(table.rows[0].cells[1], provider)
    for cell in table.rows[0].cells:
        _set_cell_borders(cell, CONFIG["border_color"])
    return table


def _title_block(doc, tb):
    s = CONFIG["sizes"]

    def centered():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        return p

    if tb.get("eyebrow"):
        _add_runs(centered(), tb["eyebrow"], bold=True, color=CONFIG["eyebrow_gray"],
                  size=s["eyebrow"], caps=True)
    p = centered()
    p.paragraph_format.space_after = Pt(3)
    _add_runs(p, tb.get("title", ""), bold=True, color=CONFIG["title_color"], size=s["title"])
    if tb.get("subtitle"):
        _add_runs(centered(), tb["subtitle"], italic=True, color=CONFIG["eyebrow_gray"], size=s["subtitle"])

    # two-column prepared-for / prepared-by, borderless
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    pf, pby = tb.get("prepared_for", {}), tb.get("prepared_by", {})
    table = doc.add_table(rows=1, cols=2)
    _no_table_borders(table)

    def party_cell(cell, party):
        cell.paragraphs[0].text = ""
        _add_runs(cell.paragraphs[0], party.get("label", ""), bold=True,
                  color=CONFIG["label_gray"], size=s["label"], caps=True)
        org = cell.add_paragraph()
        _add_runs(org, party.get("org", ""), bold=True, size=s["body"])
        if party.get("person"):
            per = cell.add_paragraph()
            _add_runs(per, party["person"], size=s["body"])

    party_cell(table.rows[0].cells[0], pf)
    party_cell(table.rows[0].cells[1], pby)

    if tb.get("date"):
        d = doc.add_paragraph()
        d.paragraph_format.space_before = Pt(8)
        _add_runs(d, tb["date"], italic=True, color=CONFIG["eyebrow_gray"], size=s["body"])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build(content, out_path):
    # Voice guard: em-dashes are banned in the proposal. Fail loudly so they get rewritten,
    # rather than shipping them or blind-swapping to a hyphen.
    hits = []

    def _scan(v, where):
        if isinstance(v, str):
            if "—" in v:
                hits.append(f"{where}: {v[:70]}")
        elif isinstance(v, dict):
            for k, x in v.items():
                _scan(x, f"{where}.{k}")
        elif isinstance(v, list):
            for i, x in enumerate(v):
                _scan(x, f"{where}[{i}]")

    _scan(content, "content")
    if hits:
        raise ValueError(
            "Em-dashes (—) are not allowed in the proposal. Rewrite each into separate "
            "sentences or use a comma, colon, or parentheses (never a hyphen), then rebuild:\n  "
            + "\n  ".join(hits)
        )

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
    _set_base_font(doc)

    if content.get("title_block"):
        _title_block(doc, content["title_block"])

    s = CONFIG["sizes"]
    for block in content.get("blocks", []):
        btype = block.get("type")
        if btype == "h1":
            _heading(doc, block["text"], size=s["h1"], color=CONFIG["navy"], space_before=14)
        elif btype == "h2":
            _heading(doc, block["text"], size=s["h2"], color=CONFIG["navy"], space_before=10)
        elif btype == "h3":
            _heading(doc, block["text"], size=s["h3"], color=CONFIG["navy"], space_before=8)
        elif btype == "p":
            _body(doc, block["text"])
        elif btype == "bullets":
            _bullets(doc, block.get("items", []))
        elif btype == "table":
            _data_table(doc, block["header"], block.get("rows", []),
                        emphasis_rows=block.get("emphasis_rows"))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        elif btype == "status_table":
            _data_table(doc, block["header"], block.get("rows", []),
                        status_col=block.get("status_col"))
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        elif btype == "signature":
            _signature_block(doc, block.get("company", {}), block.get("provider", {}))
        elif btype == "spacer":
            doc.add_paragraph()

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 3:
        print("usage: python3 build_proposal_docx.py <content.json> <out.docx>", file=sys.stderr)
        sys.exit(2)
    with open(sys.argv[1], "r", encoding="utf-8") as f:
        content = json.load(f)
    path = build(content, sys.argv[2])
    print(f"wrote {path}")


if __name__ == "__main__":
    main()
